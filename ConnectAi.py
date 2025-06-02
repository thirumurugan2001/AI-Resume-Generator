from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os 
import tempfile
from dotenv import load_dotenv
import random
import os
load_dotenv()
from helper import *
from openai import OpenAI

def generate_resume_with_gpt(name, email, phone, skills, experience, summary, linked_in, degree, university, grad_year, projects):
    try : 
        prompt = f"""
        Generate a professional resume for {name} with {experience} years of experience.
        The resume should be formatted properly and include:
        - Full Name: {name}
        - Contact Information: {email} | {phone} | {linked_in}
        - Professional Summary: {summary if summary else f'A software developer with {experience} years of experience in Python and Django.'}
        - Key Skills: {skills}
        - Education: {degree} from {university}, Graduated in {grad_year}
        - Projects: {projects}
        - Work Experience: {experience} years of relevant experience.
        The resume should be well-structured, clean, and visually appealing.
        """
        try:
            client = OpenAI(
                api_key=os.getenv("AZURE_OPENAI_KY"),
                base_url=os.getenv("GPT_URL")
            )

            response = client.chat.completions.create(
                model=os.getenv("MODEL"),
                messages=[
                    {"role": "system", "content": "You are a professional resume writer."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.8,
                max_tokens=4096
            )
            generated_resume = response.choices[0].message.content

            print("GENERATED_RESUME:", generated_resume)
        except Exception as e:
            print("Error in the generate_resume_with_gpt", str(e))
            return {
                "data": [{}],
                "statusCode":200,
                "message": "Failed to generate resume",
                "status":False
            }

        doc = Document()
        title = doc.add_heading("RESUME", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.runs[0].font.color.rgb = RGBColor(31, 78, 121)
        
        for line in generated_resume.split('\n'):
            if line.strip():
                line = line.replace("**", "").strip()
                line = line.replace("---", "").strip()
                
                if line.startswith("# ") or line.startswith("### "):
                    heading_text = line.lstrip("#").strip()
                    heading = doc.add_heading(heading_text, level=2)
                    heading.runs[0].font.color.rgb = RGBColor(31, 78, 121)
                elif line.startswith("-"):
                    doc.add_paragraph(line)
                elif "📧" in line or "📞" in line or "🔗" in line:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(line)
                    run.bold = True
                    run.font.size = Pt(10)
                else:
                    doc.add_paragraph(line)
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file_path = temp_file.name
        temp_file.close()
        doc.save(temp_file_path)
        filename = random.randint(10000, 99999)
        filename=f"generated_resume{filename}.docx"
        success, url = upload_cs_file(temp_file_path,filename)
        if success:     
            return {
                "data": [{
                    "file_url":url
                }],
                "statusCode":200,
                "message": "Resume generated successfully",
                "status":True
            }
        else :
            return {
                "data": [{}],
                "statusCode":200,
                "message": "Failed to generate resume",
                "status":False
            }
    except Exception as e :
        print("Error in the generate_resume_with_gpt", str(e))
        return {
            "message": str(e),
            "statusCode":400,
            "status":False,
            "data":{[]}
        }
